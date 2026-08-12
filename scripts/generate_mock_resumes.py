import fitz
import docx
from pathlib import Path

def generate_pdf(path: Path, text: str):
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    # Simple write lines line-by-line
    y = 50
    for line in text.split("\n"):
        page.insert_text((50, y), line)
        y += 18
    pdf_doc.save(path)
    pdf_doc.close()

def generate_docx(path: Path, text: str):
    doc = docx.Document()
    for line in text.split("\n"):
        clean = line.strip()
        if clean.startswith("Summary:") or clean.startswith("Experience:") or clean.startswith("Education:") or clean.startswith("Skills:"):
            doc.add_heading(clean, level=2)
        elif clean:
            doc.add_paragraph(clean)
    doc.save(path)

def generate_txt(path: Path, text: str):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)

def main():
    resumes_dir = Path("data/resumes")
    resumes_dir.mkdir(parents=True, exist_ok=True)
    
    candidates = [
        # Candidate 01: PDF, strong fit (Python, FastAPI, PyTorch, Docker, B.Tech, 3 years exp)
        (
            "candidate_01.pdf",
            """John Doe
Backend AI/ML Engineer
Email: john.doe@example.com | Phone: +1-555-0101

Summary:
Experienced Software Engineer specializing in AI pipelines and FastAPI backend services.

Experience:
Tech Corp - Software Engineer
3 years of experience.
- Designed and deployed FastAPI backend apps with SQL.
- Used Docker for containerizing models.

Education:
B.Tech in Computer Science from State University.

Skills: Python, FastAPI, Docker, PostgreSQL, PyTorch, scikit-learn, Git"""
        ),
        # Candidate 02: DOCX, strong fit (Python, scikit-learn, PyTorch, M.Tech, 4 years exp)
        (
            "candidate_02.docx",
            """Jane Smith
Data Scientist / ML Engineer
Email: jane.smith@example.com | Phone: +1-555-0102

Summary:
Data Scientist with extensive experience training deep learning models and writing clean Python pipelines.

Experience:
Analytics Co - ML Developer
4 years of experience in AI/ML field.
- Trained deep learning architectures using PyTorch and TensorFlow.
- Cleaned and prepared large datasets using Pandas and NumPy.

Education:
M.Tech in Artificial Intelligence

Skills: Python, PyTorch, TensorFlow, scikit-learn, Pandas, NumPy, SQL, Git"""
        ),
        # Candidate 03: TXT, moderate fit (Python, FastAPI, Docker, Bachelor, 2 years exp)
        (
            "candidate_03.txt",
            """Alex Jones
DevOps Engineer
Email: alex.jones@example.com | Phone: +1-555-0103

Summary:
Infrastructure Engineer focused on containerized workloads and CI/CD pipelines.

Experience:
CloudCorp - DevOps Analyst
2 years experience.
- Implemented Docker and Kubernetes platforms.
- Automated cloud deployments on AWS.

Education:
Bachelor of Science in Information Technology

Skills: Docker, Kubernetes, AWS, Git, Python, Bash"""
        ),
        # Candidate 04: PDF, moderate fit (TensorFlow, scikit-learn, PhD, 1.5 years exp)
        (
            "candidate_04.pdf",
            """David Miller
AI Researcher
Email: david.miller@example.com | Phone: 555-0104

Summary:
AI Researcher with strong math background and publications in deep learning.

Experience:
AI Labs - Research Assistant
1.5 years experience.
- Formulated machine learning models using TensorFlow.
- Evaluated models using scikit-learn.

Education:
PhD in Computer Science

Skills: Python, TensorFlow, scikit-learn, Git, Deep Learning"""
        ),
        # Candidate 05: DOCX, weak fit / wrong stack (Java, C++, SQL, Bachelor, 5 years exp)
        (
            "candidate_05.docx",
            """Emily Brown
Enterprise Developer
Email: emily.brown@example.com | Phone: 555-0105

Summary:
Backend developer experienced in enterprise Java application development.

Experience:
Enterprise Corp - Java Developer
5 years experience.
- Created REST APIs using Java and Spring Boot.
- Wrote complex SQL procedures for databases.

Education:
Bachelor in Computer Engineering

Skills: Java, C++, SQL, Git"""
        ),
        # Candidate 06: TXT, weak fit / no exp (Python, FastAPI, B.Tech, 0 years exp)
        (
            "candidate_06.txt",
            """Michael Green
Graduate Software Developer
Email: michael.green@example.com | Phone: 555-0106

Summary:
Fresh graduate looking for an entry level software engineering role.

Experience:
No professional experience.

Education:
B.Tech in Computer Science (Graduated 2026)

Skills: Python, FastAPI, scikit-learn, SQL"""
        ),
        # Candidate 07: PDF, strong skills / short experience (Python, PyTorch, Docker, BCA, 0.5 years exp)
        (
            "candidate_07.pdf",
            """Sophia White
ML Intern
Email: sophia.white@example.com | Phone: 555-0107

Summary:
Self-taught ML developer with strong open source contributions.

Experience:
Internship at StartUp Inc
0.5 years experience.
- Developed PyTorch pipelines and FastAPI backends.
- Used Docker and Kubernetes.

Education:
BCA in Computer Applications

Skills: Python, PyTorch, FastAPI, Docker, Kubernetes, Git"""
        ),
        # Candidate 08: DOCX, strong experience / wrong stack (Java, AWS, Docker, MCA, 6 years exp)
        (
            "candidate_08.docx",
            """Robert Black
Principal Engineer
Email: robert.black@example.com | Phone: 555-0108

Summary:
Senior systems architect with a history of building scalable web apps.

Experience:
Cloud Platforms - Principal Engineer
6 years experience.
- Automated workloads on AWS and Docker.
- Wrote backend services in Node.js and Java.

Education:
MCA (Master of Computer Applications)

Skills: Java, JavaScript, Node.js, SQL, AWS, Docker, Git"""
        ),
        # Candidate 09: TXT, strong education / weak skills (Python, C++, PhD, 1 year exp)
        (
            "candidate_09.txt",
            """William Grey
Astrophysics PhD Graduate
Email: william.grey@example.com | Phone: 555-0109

Summary:
Astrophysicist with strong scientific computing background.

Experience:
Research Fellow at Observatory
1 year experience.
- Simulated physical models using Python and C++.

Education:
PhD in Astrophysics

Skills: Python, C++, Git"""
        ),
        # Candidate 10: PDF, strong fit (Python, NLP, LLM, RAG, B.E, 2.5 years exp)
        (
            "candidate_10.pdf",
            """Olivia Taylor
NLP Developer
Email: olivia.taylor@example.com | Phone: 555-0110

Summary:
NLP developer specializing in Large Language Models (LLMs) and RAG applications.

Experience:
Cognitive Solutions - AI Engineer
2.5 years experience.
- Deployed RAG and LLM systems in production.
- Used Docker, Kubernetes, and Git.

Education:
B.E in Computer Science

Skills: Python, NLP, LLM, RAG, PyTorch, Git, Docker, Kubernetes, AWS"""
        )
    ]
    
    for filename, content in candidates:
        filepath = resumes_dir / filename
        ext = filepath.suffix.lower()
        if ext == ".pdf":
            generate_pdf(filepath, content)
        elif ext == ".docx":
            generate_docx(filepath, content)
        else:
            generate_txt(filepath, content)
        print(f"Generated synthetic candidate file: {filepath}")

if __name__ == "__main__":
    main()
