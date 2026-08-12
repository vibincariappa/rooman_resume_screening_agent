import pytest
from pathlib import Path
import fitz
import docx
from app.parsers import parse_document, normalize_text
from app.models.schemas import ResumeDocument

@pytest.fixture
def temp_dir(tmp_path):
    return tmp_path

def test_normalize_text():
    raw = "Line 1 \r\n\r\n\r\n Line 2 \r\n  \r\n\r\n\r\nLine 3"
    normalized = normalize_text(raw)
    assert normalized == "Line 1\n\nLine 2\n\nLine 3"

def test_txt_parser(temp_dir):
    txt_path = temp_dir / "candidate.txt"
    content = "John Doe\n\nPython Software Engineer\n\nSkills: Python, FastAPI."
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content)
        
    doc = parse_document(txt_path)
    assert isinstance(doc, ResumeDocument)
    assert doc.filename == "candidate.txt"
    assert doc.file_type == "txt"
    assert "John Doe" in doc.normalized_text
    assert "FastAPI" in doc.normalized_text
    assert doc.character_count == len(doc.normalized_text)
    assert doc.word_count == len(doc.normalized_text.split())

def test_pdf_parser(temp_dir):
    pdf_path = temp_dir / "candidate.pdf"
    
    # Programmatically create a valid PDF using fitz
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text((50, 50), "Jane Smith\n\nData Scientist\n\nSkills: PyTorch, scikit-learn.")
    pdf_doc.save(pdf_path)
    pdf_doc.close()
    
    doc = parse_document(pdf_path)
    assert isinstance(doc, ResumeDocument)
    assert doc.filename == "candidate.pdf"
    assert doc.file_type == "pdf"
    assert "Jane Smith" in doc.normalized_text
    assert "PyTorch" in doc.normalized_text

def test_docx_parser(temp_dir):
    docx_path = temp_dir / "candidate.docx"
    
    # Programmatically create a valid DOCX
    docx_doc = docx.Document()
    docx_doc.add_paragraph("Alice Johnson")
    docx_doc.add_paragraph("DevOps Engineer")
    table = docx_doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Skills: Docker, Kubernetes"
    docx_doc.save(docx_path)
    
    doc = parse_document(docx_path)
    assert isinstance(doc, ResumeDocument)
    assert doc.filename == "candidate.docx"
    assert doc.file_type == "docx"
    assert "Alice Johnson" in doc.normalized_text
    assert "Docker" in doc.normalized_text

def test_missing_file():
    missing_path = Path("non_existent_file.pdf")
    with pytest.raises(FileNotFoundError):
        parse_document(missing_path)

def test_unsupported_extension(temp_dir):
    invalid_path = temp_dir / "candidate.xyz"
    with open(invalid_path, "w") as f:
        f.write("Some text")
        
    with pytest.raises(ValueError, match="Unsupported file type"):
        parse_document(invalid_path)

def test_empty_txt(temp_dir):
    empty_path = temp_dir / "empty.txt"
    with open(empty_path, "w") as f:
        f.write("")
        
    with pytest.raises(ValueError, match="Empty or unreadable"):
        parse_document(empty_path)

def test_empty_pdf(temp_dir):
    pdf_path = temp_dir / "empty.pdf"
    pdf_doc = fitz.open()
    pdf_doc.new_page() # blank page
    pdf_doc.save(pdf_path)
    pdf_doc.close()
    
    with pytest.raises(ValueError, match="Empty or unreadable"):
        parse_document(pdf_path)

def test_empty_docx(temp_dir):
    docx_path = temp_dir / "empty.docx"
    docx_doc = docx.Document()
    docx_doc.save(docx_path)
    
    with pytest.raises(ValueError, match="Empty or unreadable"):
        parse_document(docx_path)

def test_corrupted_pdf(temp_dir):
    corrupt_path = temp_dir / "corrupt.pdf"
    with open(corrupt_path, "wb") as f:
        f.write(b"Not a PDF header")
        
    with pytest.raises(ValueError, match="Corrupted or invalid PDF"):
        parse_document(corrupt_path)

def test_corrupted_docx(temp_dir):
    corrupt_path = temp_dir / "corrupt.docx"
    with open(corrupt_path, "wb") as f:
        f.write(b"Not a Zip archive that DOCX requires")
        
    with pytest.raises(ValueError, match="Corrupted or invalid DOCX"):
        parse_document(corrupt_path)

def test_skills_matching():
    from app.extraction.skills import match_skills_in_text, extract_skills_categorized
    
    # Test case-insensitivity
    text_casing = "python and fastapi and SQL"
    matched = match_skills_in_text(text_casing)
    assert "Python" in matched
    assert "FastAPI" in matched
    assert "SQL" in matched
    
    # Test boundary/false positive avoidance
    text_false_positives = "Java developer. JavaScript also preferred."
    matched_false = match_skills_in_text(text_false_positives)
    assert "Java" in matched_false
    assert "JavaScript" in matched_false
    
    text_java_only = "Java developer."
    matched_java = match_skills_in_text(text_java_only)
    assert "Java" in matched_java
    assert "JavaScript" not in matched_java
    
    # Test special characters (C++, Next.js)
    text_special = "We need C++ and Next.js developers."
    matched_spec = match_skills_in_text(text_special)
    assert "C++" in matched_spec
    assert "Next.js" in matched_spec
    
    # Test sections categorization (required vs. preferred)
    jd_text = """
    Software Engineer
    
    Requirements:
    - Experience with Python and Django.
    
    Preferred Qualifications:
    - Docker and AWS knowledge is a plus.
    """
    req, pref = extract_skills_categorized(jd_text)
    assert "Python" in req
    assert "Django" in req
    assert "Docker" in pref
    assert "AWS" in pref
    assert "Python" not in pref

def test_experience_extraction():
    from app.extraction.experience import extract_minimum_experience_years
    
    assert extract_minimum_experience_years("3 years experience") == 3.0
    assert extract_minimum_experience_years("3+ years") == 3.0
    assert extract_minimum_experience_years("minimum 2 years of experience") == 2.0
    assert extract_minimum_experience_years("Required: at least 5 years") == 5.0
    assert extract_minimum_experience_years("no experience required") == 0.0

def test_education_extraction():
    from app.extraction.education import extract_education_requirements
    
    assert "B.Tech" in extract_education_requirements("Requires a B.Tech or M.Tech degree")
    assert "M.Tech" in extract_education_requirements("Requires a B.Tech or M.Tech degree")
    assert "PhD" in extract_education_requirements("PhD preferred")
    assert "Bachelor" in extract_education_requirements("Bachelor's degree in CS")

def test_job_description_extraction_end_to_end():
    from app.extraction.jd_extractor import extract_job_description
    
    jd_text = """
    Junior AI/ML Engineer
    
    Responsibilities:
    - Design and train ML models.
    - Preprocess raw datasets.
    
    Requirements:
    - Bachelor or B.Tech degree.
    - 2+ years of experience.
    - Python and PyTorch.
    
    Preferred Qualifications:
    - Docker and AWS.
    """
    
    jd = extract_job_description(jd_text)
    assert jd.title == "Junior AI/ML Engineer"
    assert jd.minimum_experience_years == 2.0
    assert "Bachelor" in jd.education_requirements
    assert "B.Tech" in jd.education_requirements
    assert "Python" in jd.required_skills
    assert "PyTorch" in jd.required_skills
    assert "Docker" in jd.preferred_skills
    assert "AWS" in jd.preferred_skills
    assert len(jd.responsibilities) == 2
    assert "Design and train ML models." in jd.responsibilities

def test_resume_name_extractor():
    from app.extraction.resume_extractor import extract_name
    
    resume_text = "John Doe\nSoftware Engineer\njohn.doe@example.com\n"
    assert extract_name(resume_text) == "John Doe"
    
    resume_text_with_header = "RESUME\n\nJane Smith\nData Scientist\n"
    assert extract_name(resume_text_with_header) == "Jane Smith"

def test_resume_contact_extractor():
    from app.extraction.resume_extractor import extract_email, extract_phone
    
    text = "Contact: test.candidate_01@domain.co.uk or call +1 (555) 0199 for info."
    assert extract_email(text) == "test.candidate_01@domain.co.uk"
    assert extract_phone(text) == "+1 (555) 0199"

def test_resume_sections_extractor():
    from app.extraction.resume_extractor import extract_work_experience_section, extract_summary_section
    
    resume = """
    Alice Johnson
    
    Summary
    Dynamic developer with experience.
    
    Experience
    Tech Corp - Software Developer
    - Coded FastAPI backend pipelines.
    
    Education
    B.Tech CS
    """
    
    summary = extract_summary_section(resume)
    assert summary == "Dynamic developer with experience."
    
    exp = extract_work_experience_section(resume)
    assert exp == "Tech Corp - Software Developer\n- Coded FastAPI backend pipelines."

def test_candidate_profile_extraction_end_to_end():
    from app.extraction.resume_extractor import extract_candidate_profile
    
    resume = """
    Bob Smith
    Email: bob.smith@work.com | Phone: 555-123-4567
    
    Summary:
    Expert data practitioner.
    
    Work Experience:
    StartUp Inc - ML Architect
    3+ years.
    - Built TensorFlow neural networks.
    
    Education:
    PhD in Math
    """
    
    profile = extract_candidate_profile(resume, "bob.txt", "bob_id")
    assert profile.candidate_id == "bob_id"
    assert profile.filename == "bob.txt"
    assert profile.name == "Bob Smith"
    assert profile.email == "bob.smith@work.com"
    assert profile.phone == "555-123-4567"
    assert "TensorFlow" in profile.skills
    assert "PhD" in profile.education
    assert profile.years_of_experience == 3.0
    assert "ML Architect" in profile.work_experience
    assert "Expert data practitioner." in profile.summary

def test_similarity_range_and_identical():
    from app.matching.embeddings import generate_embedding
    from app.matching.similarity import calculate_similarity
    
    text = "We are seeking a senior python engineer with experience in cloud environments."
    emb_a = generate_embedding(text)
    emb_b = generate_embedding(text)
    
    sim = calculate_similarity(emb_a, emb_b)
    # Identical text must yield 100.0
    assert sim == 100.0
    
    # Test ranges
    text_diff = "A completely different sentence discussing football and cooking recipes."
    emb_diff = generate_embedding(text_diff)
    sim_diff = calculate_similarity(emb_a, emb_diff)
    assert 0.0 <= sim_diff <= 100.0
    # Unrelated texts should be significantly less than 100.0
    assert sim_diff < 50.0

def test_model_loading_caching():
    from app.matching.embeddings import get_embedding_model
    
    model_1 = get_embedding_model("all-MiniLM-L6-v2")
    model_2 = get_embedding_model("all-MiniLM-L6-v2")
    
    # The references must be identical (same object)
    assert model_1 is model_2

def test_multiple_candidates_similarity():
    from app.matching.embeddings import generate_embedding
    from app.matching.similarity import calculate_similarity
    
    jd = "Seeking a machine learning expert in PyTorch and TensorFlow."
    c1 = "Deep learning engineer with PyTorch and TensorFlow expertise."
    c2 = "Front-end developer experienced in React and Javascript."
    
    emb_jd = generate_embedding(jd)
    emb_c1 = generate_embedding(c1)
    emb_c2 = generate_embedding(c2)
    
    sim_1 = calculate_similarity(emb_jd, emb_c1)
    sim_2 = calculate_similarity(emb_jd, emb_c2)
    
    # Candidate 1 must be more similar to the JD than candidate 2
    assert sim_1 > sim_2

def test_candidate_scoring_perfect():
    from app.models.schemas import CandidateProfile, JobDescription
    from app.matching.scoring import calculate_candidate_score
    
    jd = JobDescription(
        title="Python Engineer",
        required_skills=["Python", "FastAPI"],
        preferred_skills=["Docker"],
        minimum_experience_years=3.0,
        education_requirements=["B.Tech"],
        responsibilities=["Write backend APIs"],
        raw_text="Job Description text"
    )
    
    profile = CandidateProfile(
        candidate_id="perfect",
        filename="perfect.txt",
        name="Perfect Candidate",
        skills=["Python", "FastAPI", "Docker"],
        education=["B.Tech in Computer Science"],
        years_of_experience=4.0,
        raw_text="Resume text"
    )
    
    breakdown = calculate_candidate_score(profile, jd, semantic_score=100.0)
    
    assert breakdown.skills_score == 100.0
    assert breakdown.experience_score == 100.0
    assert breakdown.education_score == 100.0
    assert breakdown.semantic_score == 100.0
    assert breakdown.final_score == 100.0

def test_candidate_scoring_missing_skills():
    from app.models.schemas import CandidateProfile, JobDescription
    from app.matching.scoring import calculate_candidate_score
    
    jd = JobDescription(
        title="Python Engineer",
        required_skills=["Python", "FastAPI"],
        preferred_skills=[],
        minimum_experience_years=0.0,
        education_requirements=[],
        responsibilities=[],
        raw_text="Job Description text"
    )
    
    profile = CandidateProfile(
        candidate_id="missing_skills",
        filename="c1.txt",
        name="Candidate A",
        skills=["Python"],
        education=[],
        years_of_experience=0.0,
        raw_text="Resume text"
    )
    
    breakdown = calculate_candidate_score(profile, jd, semantic_score=80.0)
    assert breakdown.skills_score == 60.0

def test_candidate_scoring_insufficient_exp():
    from app.models.schemas import CandidateProfile, JobDescription
    from app.matching.scoring import calculate_candidate_score
    
    jd = JobDescription(
        title="Python Engineer",
        required_skills=[],
        preferred_skills=[],
        minimum_experience_years=4.0,
        education_requirements=[],
        responsibilities=[],
        raw_text="Job Description text"
    )
    
    profile = CandidateProfile(
        candidate_id="low_exp",
        filename="c2.txt",
        name="Candidate B",
        skills=[],
        education=[],
        years_of_experience=2.0,
        raw_text="Resume text"
    )
    
    breakdown = calculate_candidate_score(profile, jd, semantic_score=80.0)
    assert breakdown.experience_score == 50.0

def test_candidate_scoring_missing_education():
    from app.models.schemas import CandidateProfile, JobDescription
    from app.matching.scoring import calculate_candidate_score
    
    jd = JobDescription(
        title="Python Engineer",
        required_skills=[],
        preferred_skills=[],
        minimum_experience_years=0.0,
        education_requirements=["PhD"],
        responsibilities=[],
        raw_text="Job Description text"
    )
    
    profile_other = CandidateProfile(
        candidate_id="btech",
        filename="c3.txt",
        skills=[],
        education=["B.Tech"],
        years_of_experience=0.0,
        raw_text="Resume text"
    )
    
    profile_none = CandidateProfile(
        candidate_id="none",
        filename="c4.txt",
        skills=[],
        education=[],
        years_of_experience=0.0,
        raw_text="Resume text"
    )
    
    br_other = calculate_candidate_score(profile_other, jd, semantic_score=80.0)
    br_none = calculate_candidate_score(profile_none, jd, semantic_score=80.0)
    
    assert br_other.education_score == 50.0
    assert br_none.education_score == 0.0

def test_candidate_scoring_no_requirements():
    from app.models.schemas import CandidateProfile, JobDescription
    from app.matching.scoring import calculate_candidate_score
    
    jd = JobDescription(
        title="Generic Role",
        required_skills=[],
        preferred_skills=[],
        minimum_experience_years=0.0,
        education_requirements=[],
        responsibilities=[],
        raw_text="JD text"
    )
    
    profile = CandidateProfile(
        candidate_id="candidate",
        filename="c.txt",
        skills=[],
        education=[],
        years_of_experience=0.0,
        raw_text="Resume text"
    )
    
    breakdown = calculate_candidate_score(profile, jd, semantic_score=80.0)
    
    assert breakdown.skills_score == 100.0
    assert breakdown.experience_score == 100.0
    assert breakdown.education_score == 100.0
    assert breakdown.final_score == 94.0



