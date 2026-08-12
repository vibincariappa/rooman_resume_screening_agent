import docx
from pathlib import Path

def parse_docx(file_path: Path) -> str:
    """
    Parses a DOCX file and returns its raw text contents.
    Raises FileNotFoundError if file does not exist.
    Raises ValueError if document is corrupted or empty.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
        
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        raise ValueError(f"Corrupted or invalid DOCX file: {e}")
        
    text_parts = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)
            
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_parts.append(cell.text)
                    
    raw_text = "\n".join(text_parts)
    
    if not raw_text.strip():
        raise ValueError("Empty or unreadable DOCX document.")
        
    return raw_text
